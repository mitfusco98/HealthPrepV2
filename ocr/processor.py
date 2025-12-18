"""
Tesseract integration and text cleanup for medical documents
Supports PDF, images, Word documents (.docx, .doc), RTF, HTML, EML, and plain text

Performance optimizations:
- Machine-readable PDF detection: Extracts embedded text before OCR (20-30% faster)
- Multi-library PDF extraction: PyMuPDF → pdfminer.six → OCR fallback chain
- Per-page hybrid processing: Only OCR pages that lack embedded text
- Configurable parallel workers via OCR_MAX_WORKERS environment variable
- Batch processing with isolated database sessions for thread safety
"""
import os
import subprocess
import tempfile
from PIL import Image
import pdf2image
from app import db
from models import Document
from .phi_filter import PHIFilter
import logging
from datetime import datetime
import email
from email import policy

try:
    import fitz  # PyMuPDF for embedded text extraction
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from striprtf.striprtf import rtf_to_text
    STRIPRTF_AVAILABLE = True
except ImportError:
    STRIPRTF_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

# Configurable worker count - defaults to 4, can be overridden via environment
MIN_TEXT_LENGTH_FOR_SKIP_OCR = 100  # Minimum chars to consider PDF machine-readable

def get_ocr_max_workers():
    """
    Get the maximum number of OCR workers from environment or auto-detect.
    
    Priority:
    1. OCR_MAX_WORKERS environment variable
    2. Auto-detect based on CPU cores (cores - 1, minimum 2)
    3. Fallback to 4 workers
    """
    env_workers = os.environ.get('OCR_MAX_WORKERS')
    if env_workers:
        try:
            workers = int(env_workers)
            if workers > 0:
                return workers
        except ValueError:
            pass
    
    # Auto-detect based on CPU cores
    try:
        import multiprocessing
        cpu_count = multiprocessing.cpu_count()
        # Use cores - 1 to leave headroom, minimum 2
        return max(2, cpu_count - 1)
    except Exception:
        pass
    
    return 4  # Fallback default

class OCRProcessor:
    """Handles OCR processing of medical documents using Tesseract"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.phi_filter = PHIFilter()

        # Tesseract configuration for medical documents
        self.tesseract_config = '--oem 3 --psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,()[]{}:;/\\-+=%$@#!?"\' \n\t'

    def process_document(self, document_id, skip_screening_update=False, secure_delete_original=True):
        """Process a document with OCR and PHI filtering
        
        HIPAA COMPLIANCE: After successful text extraction, the original file is
        securely deleted (overwritten with random data before unlinking) to ensure
        PHI cannot be recovered from disk.
        
        Args:
            document_id: ID of document to process
            skip_screening_update: If True, skip the synchronous screening update.
                                  Use this when processing as part of a batch where
                                  screening updates will be done after all OCR completes.
            secure_delete_original: If True (default), securely delete the original file
                                   after successful text extraction for HIPAA compliance.
        """
        document = Document.query.get(document_id)
        if not document:
            self.logger.error(f"Document {document_id} not found")
            return False

        original_file_path = document.file_path
        
        try:
            # Extract text using OCR
            ocr_text, confidence = self._extract_text(document.file_path)

            if ocr_text:
                # Apply PHI filtering if enabled
                filtered_text = self.phi_filter.filter_phi(ocr_text)

                # Update document record
                document.ocr_text = filtered_text
                document.content = filtered_text  # Also update content field for backward compatibility
                document.ocr_confidence = confidence
                document.phi_filtered = True
                document.processed_at = datetime.utcnow()

                db.session.commit()

                self.logger.info(f"Successfully processed document {document_id} with confidence {confidence:.2f}")
                
                # HIPAA: Securely delete original file after successful extraction
                if secure_delete_original and original_file_path:
                    self._secure_delete_original(document, original_file_path)

                # Trigger screening engine update (unless skipped for batch processing)
                if not skip_screening_update:
                    from core.engine import ScreeningEngine
                    engine = ScreeningEngine()
                    engine.process_new_document(document_id)

                return True
            else:
                self.logger.warning(f"No text extracted from document {document_id}")
                return False

        except Exception as e:
            self.logger.error(f"Error processing document {document_id}: {str(e)}")
            return False
    
    def _secure_delete_original(self, document, file_path):
        """Securely delete original file after text extraction for HIPAA compliance.
        
        This overwrites the file with random data before deletion to prevent
        PHI recovery from disk.
        """
        try:
            from utils.secure_delete import secure_delete_file, audit_log_deletion
            
            if secure_delete_file(file_path):
                # Update document record to reflect disposal
                document.file_path = None
                document.file_disposed = True
                document.file_disposed_at = datetime.utcnow()
                db.session.commit()
                
                # Log for HIPAA audit trail
                audit_log_deletion(
                    file_path=file_path,
                    file_type='uploaded_document',
                    patient_id=document.patient_id,
                    org_id=document.org_id
                )
                
                self.logger.info(f"HIPAA: Securely deleted original file for document {document.id}")
            else:
                self.logger.warning(f"Failed to securely delete original file for document {document.id}")
                
        except Exception as e:
            self.logger.error(f"Error during secure deletion of original file: {e}")

    def _extract_text(self, file_path):
        """Extract text from document using appropriate method"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document file not found: {file_path}")

        file_ext = os.path.splitext(file_path)[1].lower()

        try:
            if file_ext == '.pdf':
                return self._process_pdf(file_path)
            elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._process_image(file_path)
            elif file_ext == '.docx':
                return self._process_docx(file_path)
            elif file_ext == '.doc':
                return self._process_doc(file_path)
            elif file_ext == '.rtf':
                return self._process_rtf(file_path)
            elif file_ext in ['.html', '.htm']:
                return self._process_html(file_path)
            elif file_ext == '.eml':
                return self._process_eml(file_path)
            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read(), 1.0
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")

        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None, 0.0

    def _extract_embedded_pdf_text(self, pdf_path):
        """
        Extract embedded text from PDF using PyMuPDF.
        Returns (text, is_machine_readable) tuple.
        
        Machine-readable PDFs have embedded text layers (not scanned images).
        This is ~100x faster than OCR and should be attempted first.
        """
        if not PYMUPDF_AVAILABLE:
            return None, False
        
        try:
            doc = fitz.open(pdf_path)
            all_text = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text")
                if text and text.strip():
                    all_text.append(text.strip())
            
            doc.close()
            
            combined_text = '\n'.join(all_text)
            
            # Check if we have substantial text (not just headers/footers)
            is_machine_readable = len(combined_text) >= MIN_TEXT_LENGTH_FOR_SKIP_OCR
            
            if is_machine_readable:
                self.logger.info(
                    f"PDF is machine-readable: {len(combined_text)} chars extracted, skipping OCR"
                )
            
            return combined_text if is_machine_readable else None, is_machine_readable
            
        except Exception as e:
            self.logger.warning(f"Error extracting embedded PDF text: {str(e)}")
            return None, False

    def _extract_pdf_text_pdfminer(self, pdf_path):
        """
        Fallback PDF text extraction using pdfminer.six.
        Used when PyMuPDF returns insufficient text.
        """
        if not PDFMINER_AVAILABLE:
            return None, False
        
        try:
            text = pdfminer_extract_text(pdf_path)
            if text and len(text.strip()) >= MIN_TEXT_LENGTH_FOR_SKIP_OCR:
                self.logger.info(
                    f"pdfminer.six extracted {len(text)} chars from PDF, skipping OCR"
                )
                return text.strip(), True
            return None, False
        except Exception as e:
            self.logger.warning(f"pdfminer.six extraction failed: {str(e)}")
            return None, False

    def _process_pdf(self, pdf_path):
        """
        Process PDF document and extract text.
        
        Extraction chain (in order of preference):
        1. PyMuPDF embedded text extraction (fastest)
        2. pdfminer.six text extraction (catches edge cases)
        3. OCR via Tesseract (slowest, for scanned documents)
        
        This reduces processing time by 20-30% for typical document sets.
        """
        # First, try to extract embedded text with PyMuPDF (fastest)
        embedded_text, is_machine_readable = self._extract_embedded_pdf_text(pdf_path)
        
        if is_machine_readable and embedded_text:
            # Machine-readable PDF - no OCR needed, confidence is 1.0
            return embedded_text, 1.0
        
        # Second, try pdfminer.six as fallback (catches some PDFs PyMuPDF misses)
        pdfminer_text, pdfminer_success = self._extract_pdf_text_pdfminer(pdf_path)
        
        if pdfminer_success and pdfminer_text:
            return pdfminer_text, 0.95  # Slightly lower confidence than PyMuPDF
        
        # Fall back to hybrid per-page processing
        # For each page: use embedded text if available, otherwise OCR
        self.logger.info(f"PDF requires hybrid per-page processing: {pdf_path}")
        
        return self._process_pdf_hybrid(pdf_path)

    def _process_pdf_hybrid(self, pdf_path):
        """
        Hybrid per-page PDF processing.
        
        For each page:
        - If the page has embedded text (>50 chars), use it directly
        - If not, OCR that specific page
        
        This optimizes mixed PDFs where some pages are scanned and some are digital.
        """
        all_text = []
        confidences = []
        pages_extracted = 0
        pages_ocred = 0
        
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                if PYMUPDF_AVAILABLE:
                    doc = fitz.open(pdf_path)
                    num_pages = len(doc)
                    
                    images = pdf2image.convert_from_path(pdf_path)
                    
                    for i in range(num_pages):
                        page = doc[i]
                        page_text = page.get_text("text").strip()
                        
                        if len(page_text) >= 50:
                            all_text.append(page_text)
                            confidences.append(1.0)
                            pages_extracted += 1
                        else:
                            if i < len(images):
                                image_path = os.path.join(temp_dir, f"page_{i}.png")
                                images[i].save(image_path, 'PNG')
                                text, confidence = self._process_image(image_path)
                                if text:
                                    all_text.append(text)
                                    confidences.append(confidence)
                                    pages_ocred += 1
                    
                    doc.close()
                else:
                    images = pdf2image.convert_from_path(pdf_path)
                    for i, image in enumerate(images):
                        image_path = os.path.join(temp_dir, f"page_{i}.png")
                        image.save(image_path, 'PNG')
                        text, confidence = self._process_image(image_path)
                        if text:
                            all_text.append(text)
                            confidences.append(confidence)
                            pages_ocred += 1
                
                self.logger.info(
                    f"Hybrid PDF processing: {pages_extracted} pages extracted, "
                    f"{pages_ocred} pages OCR'd"
                )
                
                combined_text = '\n'.join(all_text)
                avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
                
                return combined_text, avg_confidence
                
            except Exception as e:
                self.logger.error(f"Error in hybrid PDF processing {pdf_path}: {str(e)}")
                return None, 0.0

    def _process_image(self, image_path):
        """Process image and extract text using Tesseract"""
        try:
            # Preprocess image for better OCR results
            processed_image_path = self._preprocess_image(image_path)

            # Run Tesseract OCR
            cmd = [
                'tesseract',
                processed_image_path,
                'stdout',
                '--oem', '3',
                '--psm', '6'
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

            if result.returncode == 0:
                text = result.stdout.strip()

                # Get confidence score
                confidence = self._calculate_confidence(processed_image_path)

                return text, confidence
            else:
                self.logger.error(f"Tesseract failed with return code {result.returncode}: {result.stderr}")
                return None, 0.0

        except subprocess.TimeoutExpired:
            self.logger.error(f"Tesseract timeout processing {image_path}")
            return None, 0.0
        except Exception as e:
            self.logger.error(f"Error processing image {image_path}: {str(e)}")
            return None, 0.0

    def _preprocess_image(self, image_path):
        """Preprocess image to improve OCR accuracy"""
        try:
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                temp_path = temp_file.name

            # Open and preprocess image
            with Image.open(image_path) as image:
                # Convert to grayscale
                if image.mode != 'L':
                    image = image.convert('L')

                # Enhance contrast and resize if needed
                from PIL import ImageEnhance

                # Enhance contrast
                enhancer = ImageEnhance.Contrast(image)
                image = enhancer.enhance(1.5)

                # Resize if image is too small
                width, height = image.size
                if width < 800 or height < 600:
                    scale_factor = max(800 / width, 600 / height)
                    new_width = int(width * scale_factor)
                    new_height = int(height * scale_factor)
                    image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)

                # Save preprocessed image
                image.save(temp_path, 'PNG')

            return temp_path

        except Exception as e:
            self.logger.error(f"Error preprocessing image {image_path}: {str(e)}")
            return image_path  # Return original if preprocessing fails

    def _calculate_confidence(self, image_path):
        """Calculate OCR confidence score"""
        try:
            cmd = [
                'tesseract',
                image_path,
                'stdout',
                '--oem', '3',
                '--psm', '6',
                'tsv'
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

            if result.returncode == 0:
                lines = result.stdout.strip().split('\n')
                confidences = []

                for line in lines[1:]:  # Skip header
                    parts = line.split('\t')
                    if len(parts) >= 11 and parts[10].strip():  # Confidence is in column 10
                        try:
                            conf = float(parts[10])
                            if conf > 0:  # Only include positive confidences
                                confidences.append(conf)
                        except ValueError:
                            continue

                if confidences:
                    return sum(confidences) / len(confidences) / 100.0  # Convert to 0-1 scale
                else:
                    return 0.5  # Default confidence if no data
            else:
                return 0.5  # Default confidence on error

        except Exception as e:
            self.logger.error(f"Error calculating confidence for {image_path}: {str(e)}")
            return 0.5  # Default confidence on error

    def _process_docx(self, docx_path):
        """
        Extract text from modern Word documents (.docx)
        
        Uses python-docx library for direct text extraction.
        """
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx library not available for .docx processing")
            return None, 0.0
        
        try:
            doc = DocxDocument(docx_path)  # type: ignore
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            if text_parts:
                combined_text = '\n'.join(text_parts)
                self.logger.info(f"Successfully extracted {len(combined_text)} characters from DOCX")
                return combined_text, 1.0
            else:
                self.logger.warning("No text found in DOCX document")
                return None, 0.0
                
        except Exception as e:
            self.logger.error(f"Error processing DOCX {docx_path}: {str(e)}")
            return None, 0.0

    def _process_doc(self, doc_path):
        """
        Extract text from legacy Word documents (.doc)
        
        Attempts antiword, catdoc, or LibreOffice conversion.
        """
        # Method 1: Try antiword
        try:
            result = subprocess.run(
                ['antiword', doc_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
                self.logger.info(f"Successfully extracted {len(text)} characters from DOC using antiword")
                return text, 0.95
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        except Exception as e:
            self.logger.warning(f"antiword extraction failed: {str(e)}")
        
        # Method 2: Try catdoc
        try:
            result = subprocess.run(
                ['catdoc', '-w', doc_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
                self.logger.info(f"Successfully extracted {len(text)} characters from DOC using catdoc")
                return text, 0.9
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        except Exception as e:
            self.logger.warning(f"catdoc extraction failed: {str(e)}")
        
        # Method 3: Try LibreOffice conversion to PDF then OCR
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', temp_dir, doc_path],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                if result.returncode == 0:
                    pdf_name = os.path.splitext(os.path.basename(doc_path))[0] + '.pdf'
                    pdf_path = os.path.join(temp_dir, pdf_name)
                    
                    if os.path.exists(pdf_path):
                        text, confidence = self._process_pdf(pdf_path)
                        if text:
                            self.logger.info(f"Successfully extracted text from DOC via PDF conversion")
                            return text, confidence * 0.9
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        except Exception as e:
            self.logger.warning(f"LibreOffice conversion failed: {str(e)}")
        
        self.logger.error(f"Unable to extract text from DOC file: {doc_path}")
        return None, 0.0

    def _process_rtf(self, rtf_path):
        """
        Extract text from RTF (Rich Text Format) files.
        Common format for medical referrals and legacy documents.
        """
        if not STRIPRTF_AVAILABLE:
            self.logger.error("striprtf library not available for RTF processing")
            return None, 0.0
        
        try:
            with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            
            text = rtf_to_text(rtf_content)
            
            if text and text.strip():
                cleaned_text = text.strip()
                self.logger.info(f"Successfully extracted {len(cleaned_text)} characters from RTF")
                return cleaned_text, 1.0
            else:
                self.logger.warning("No text found in RTF document")
                return None, 0.0
                
        except Exception as e:
            self.logger.error(f"Error processing RTF {rtf_path}: {str(e)}")
            return None, 0.0

    def _process_html(self, html_path):
        """
        Extract text from HTML files.
        Common format for patient portal exports and email attachments.
        """
        if not BS4_AVAILABLE:
            self.logger.error("BeautifulSoup library not available for HTML processing")
            return None, 0.0
        
        try:
            with open(html_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            for script in soup(["script", "style", "meta", "link", "noscript"]):
                script.decompose()
            
            text = soup.get_text(separator='\n', strip=True)
            
            if text:
                self.logger.info(f"Successfully extracted {len(text)} characters from HTML")
                return text, 1.0
            else:
                self.logger.warning("No text found in HTML document")
                return None, 0.0
                
        except Exception as e:
            self.logger.error(f"Error processing HTML {html_path}: {str(e)}")
            return None, 0.0

    def _process_eml(self, eml_path):
        """
        Extract text from EML (email) files.
        Common format for forwarded medical records and referrals.
        """
        try:
            with open(eml_path, 'rb') as f:
                msg = email.message_from_binary_file(f, policy=policy.default)
            
            text_parts = []
            
            if msg['subject']:
                text_parts.append(f"Subject: {msg['subject']}")
            if msg['from']:
                text_parts.append(f"From: {msg['from']}")
            if msg['date']:
                text_parts.append(f"Date: {msg['date']}")
            
            text_parts.append("")
            
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    if content_type == 'text/plain':
                        payload = part.get_payload(decode=True)
                        if payload:
                            charset = part.get_content_charset() or 'utf-8'
                            text_parts.append(payload.decode(charset, errors='ignore'))
                    elif content_type == 'text/html' and BS4_AVAILABLE:
                        payload = part.get_payload(decode=True)
                        if payload:
                            charset = part.get_content_charset() or 'utf-8'
                            html_content = payload.decode(charset, errors='ignore')
                            soup = BeautifulSoup(html_content, 'html.parser')
                            for script in soup(["script", "style"]):
                                script.decompose()
                            text_parts.append(soup.get_text(separator='\n', strip=True))
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    charset = msg.get_content_charset() or 'utf-8'
                    text_parts.append(payload.decode(charset, errors='ignore'))
            
            combined_text = '\n'.join(text_parts).strip()
            
            if combined_text:
                self.logger.info(f"Successfully extracted {len(combined_text)} characters from EML")
                return combined_text, 1.0
            else:
                self.logger.warning("No text found in EML document")
                return None, 0.0
                
        except Exception as e:
            self.logger.error(f"Error processing EML {eml_path}: {str(e)}")
            return None, 0.0

    def reprocess_document(self, document_id):
        """Reprocess an existing document"""
        return self.process_document(document_id)

    def get_processing_stats(self):
        """Get OCR processing statistics"""
        total_docs = Document.query.count()
        processed_docs = Document.query.filter(Document.ocr_text.isnot(None)).count()

        # Calculate average confidence
        avg_confidence = db.session.query(db.func.avg(Document.ocr_confidence)).scalar() or 0.0

        # Count low confidence documents
        low_confidence_docs = Document.query.filter(Document.ocr_confidence < 0.6).count()

        return {
            'total_documents': total_docs,
            'processed_documents': processed_docs,
            'processing_rate': (processed_docs / total_docs * 100) if total_docs > 0 else 0,
            'average_confidence': avg_confidence,
            'low_confidence_documents': low_confidence_docs
        }

    def cleanup_temp_files(self):
        """Clean up any temporary files created during processing with secure deletion.
        
        HIPAA COMPLIANCE: Uses secure deletion (overwrite before unlink) to prevent
        PHI recovery from temp files created during OCR processing.
        """
        from utils.secure_delete import secure_delete_file
        temp_dir = tempfile.gettempdir()
        deleted_count = 0

        try:
            for filename in os.listdir(temp_dir):
                if filename.startswith('tesseract_') or filename.startswith('ocr_') or filename.startswith('healthprep_'):
                    file_path = os.path.join(temp_dir, filename)
                    if os.path.isfile(file_path):
                        if secure_delete_file(file_path):
                            deleted_count += 1
            
            if deleted_count > 0:
                self.logger.info(f"HIPAA: Securely deleted {deleted_count} temp files")

        except Exception as e:
            self.logger.warning(f"Error cleaning up temp files: {str(e)}")

    def process_documents_batch(self, document_ids, max_workers=None, progress_callback=None):
        """
        Process multiple documents in parallel using ThreadPoolExecutor.
        Each thread gets its own Flask app context and session for safe database access.
        
        Args:
            document_ids: List of document IDs to process
            max_workers: Maximum number of parallel workers (None = auto-detect from 
                        OCR_MAX_WORKERS env var or CPU cores)
            progress_callback: Optional callback function(processed, total, current_doc_id)
        
        Returns:
            Dict with results summary
        """
        # Use configurable worker count
        if max_workers is None:
            max_workers = get_ocr_max_workers()
        from concurrent.futures import ThreadPoolExecutor, as_completed
        from app import app, db
        from sqlalchemy.orm import scoped_session, sessionmaker
        import time
        
        results = {
            'total': len(document_ids),
            'successful': [],
            'failed': [],
            'start_time': time.time()
        }
        
        if not document_ids:
            return results
        
        self.logger.info(f"Starting parallel OCR processing of {len(document_ids)} documents with {max_workers} workers")
        
        def process_single_with_isolated_session(doc_id):
            """Worker function with isolated session for thread-safe database access"""
            with app.app_context():
                try:
                    thread_session = scoped_session(sessionmaker(bind=db.engine))
                    try:
                        success = self._process_document_with_session(doc_id, thread_session)
                        return (doc_id, success, None)
                    finally:
                        thread_session.remove()
                except Exception as e:
                    return (doc_id, False, str(e))
        
        processed_count = 0
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_doc = {executor.submit(process_single_with_isolated_session, doc_id): doc_id for doc_id in document_ids}
            
            for future in as_completed(future_to_doc):
                doc_id = future_to_doc[future]
                processed_count += 1
                
                try:
                    doc_id, success, error = future.result()
                    
                    if success:
                        results['successful'].append(doc_id)
                    else:
                        results['failed'].append({
                            'document_id': doc_id,
                            'error': error or 'Processing failed'
                        })
                        
                except Exception as e:
                    results['failed'].append({
                        'document_id': doc_id,
                        'error': str(e)
                    })
                
                if progress_callback:
                    try:
                        progress_callback(processed_count, len(document_ids), doc_id)
                    except Exception:
                        pass
        
        results['end_time'] = time.time()
        results['duration_seconds'] = results['end_time'] - results['start_time']
        results['docs_per_second'] = len(document_ids) / results['duration_seconds'] if results['duration_seconds'] > 0 else 0
        
        self.logger.info(
            f"Parallel OCR complete: {len(results['successful'])}/{len(document_ids)} successful, "
            f"{results['docs_per_second']:.2f} docs/sec"
        )
        
        return results
    
    def _process_document_with_session(self, document_id, session):
        """
        Internal document processing with explicit session for thread safety.
        Used by parallel batch processing.
        """
        document = session.query(Document).get(document_id)
        if not document:
            self.logger.error(f"Document {document_id} not found")
            return False

        try:
            ocr_text, confidence = self._extract_text(document.file_path)

            if ocr_text:
                filtered_text = self.phi_filter.filter_phi(ocr_text)
                document.ocr_text = filtered_text
                document.content = filtered_text
                document.ocr_confidence = confidence
                document.phi_filtered = True
                document.processed_at = datetime.utcnow()
                session.commit()
                self.logger.info(f"Successfully processed document {document_id} with confidence {confidence:.2f}")
                return True
            else:
                self.logger.warning(f"No text extracted from document {document_id}")
                return False

        except Exception as e:
            self.logger.error(f"Error processing document {document_id}: {str(e)}")
            session.rollback()
            return False
    
    def process_documents_batch_with_screening_update(self, document_ids, max_workers=None, progress_callback=None):
        """
        Process documents in parallel and trigger batch screening updates after completion.
        This is the preferred method for bulk document intake.
        
        Uses batched screening refresh instead of per-document updates for efficiency.
        
        Args:
            max_workers: None = auto-detect from OCR_MAX_WORKERS env var or CPU cores
        """
        from app import app
        
        # max_workers=None will trigger auto-detection in process_documents_batch
        results = self.process_documents_batch(document_ids, max_workers, progress_callback)
        
        if results['successful']:
            with app.app_context():
                try:
                    # Get unique patient IDs from processed documents
                    processed_patient_ids = set()
                    org_id = None
                    
                    for doc_id in results['successful']:
                        doc = Document.query.get(doc_id)
                        if doc:
                            processed_patient_ids.add(doc.patient_id)
                            if org_id is None:
                                org_id = doc.org_id
                    
                    if org_id and processed_patient_ids:
                        # Use batch screening refresh for all affected patients at once
                        # This is much more efficient than per-document updates
                        from services.screening_refresh_service import ScreeningRefreshService
                        
                        refresh_service = ScreeningRefreshService(org_id)
                        refresh_result = refresh_service.refresh_screenings({
                            'force_refresh': True,
                            'patient_filter': {'patient_ids': list(processed_patient_ids)},
                            'max_patients': len(processed_patient_ids)
                        })
                        
                        if refresh_result.get('success'):
                            self.logger.info(
                                f"Batch screening refresh completed: "
                                f"{refresh_result['stats']['screenings_updated']} screenings updated "
                                f"across {len(processed_patient_ids)} patients"
                            )
                        else:
                            self.logger.warning(f"Batch screening refresh had issues: {refresh_result.get('error')}")
                    
                except Exception as e:
                    self.logger.error(f"Error updating screenings after batch: {str(e)}")
        
        return results

    def process_pdf_pages_parallel(self, pdf_path, max_workers=None):
        """
        Process PDF pages in parallel for faster OCR.
        
        Args:
            pdf_path: Path to PDF file
            max_workers: None = auto-detect from OCR_MAX_WORKERS env var or CPU cores
        
        Returns:
            Tuple of (combined_text, average_confidence)
        """
        from concurrent.futures import ThreadPoolExecutor, as_completed
        import time
        
        # Use configurable worker count
        if max_workers is None:
            max_workers = get_ocr_max_workers()
        
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                images = pdf2image.convert_from_path(pdf_path)
                
                if len(images) <= 2:
                    return self._process_pdf(pdf_path)
                
                self.logger.info(f"Processing {len(images)} PDF pages in parallel")
                
                image_paths = []
                for i, image in enumerate(images):
                    image_path = os.path.join(temp_dir, f"page_{i}.png")
                    image.save(image_path, 'PNG')
                    image_paths.append((i, image_path))
                
                page_results = {}
                
                def process_page(args):
                    page_num, image_path = args
                    text, confidence = self._process_image(image_path)
                    return (page_num, text, confidence)
                
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = {executor.submit(process_page, args): args[0] for args in image_paths}
                    
                    for future in as_completed(futures):
                        try:
                            page_num, text, confidence = future.result()
                            page_results[page_num] = (text, confidence)
                        except Exception as e:
                            page_num = futures[future]
                            self.logger.error(f"Error processing page {page_num}: {str(e)}")
                            page_results[page_num] = ('', 0.0)
                
                all_text = []
                confidences = []
                for i in range(len(images)):
                    text, confidence = page_results.get(i, ('', 0.0))
                    if text:
                        all_text.append(text)
                        confidences.append(confidence)
                
                combined_text = '\n'.join(all_text)
                avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
                
                return combined_text, avg_confidence
                
            except Exception as e:
                self.logger.error(f"Error in parallel PDF processing: {str(e)}")
                return self._process_pdf(pdf_path)