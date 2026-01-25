"""
Enhanced Document Processor for EMR Integration
Processes Epic FHIR documents with OCR, keyword matching, and screening detection
Supports PDF, images, Word documents (.docx, .doc), and HTML
"""

import os
import logging
import tempfile
import base64
import re
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
from PIL import Image
import pdf2image
import pytesseract

from models import db, FHIRDocument
from ocr.processor import OCRProcessor, get_ocr_timeout_seconds
from ocr.phi_filter import PHIFilter
from core.fuzzy_detection import FuzzyDetectionEngine
from utils.document_audit import DocumentAuditLogger
from concurrent.futures import ThreadPoolExecutor, wait, FIRST_COMPLETED


def get_max_document_pages() -> int:
    """
    Get the maximum number of pages allowed for document OCR processing.
    
    This is a COST CONTROL MECHANISM to prevent runaway compute costs from
    processing extremely large documents (e.g., 100+ page medical histories).
    
    At $300/month/provider pricing with HITRUST i2 compliance debt, it's critical
    to cap the compute time spent on any single document.
    
    Priority:
    1. MAX_DOCUMENT_PAGES environment variable
    2. Fallback to 20 pages (reasonable for most screening-relevant documents)
    
    Documents exceeding this limit are flagged as skipped_oversized in the database
    and logged for audit purposes.
    """
    env_limit = os.environ.get('MAX_DOCUMENT_PAGES')
    if env_limit:
        try:
            limit = int(env_limit)
            if limit > 0:
                return limit
        except ValueError:
            pass
    
    return 20  # Default: 20 pages (reasonable for most screening docs)

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Enhanced document processor for EMR integration
    Handles Epic FHIR documents with screening-specific processing
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Initialize components
        self.ocr_processor = OCRProcessor()
        self.phi_filter = PHIFilter()
        self.fuzzy_engine = FuzzyDetectionEngine()
        
        # Screening-specific document types
        self.screening_document_types = {
            'mammography': ['mammogram', 'mammography', 'breast imaging', 'tomosynthesis'],
            'colonoscopy': ['colonoscopy', 'sigmoidoscopy', 'colon', 'colorectal'],
            'cervical': ['pap smear', 'pap test', 'cervical', 'cytology'],
            'bone_density': ['dexa', 'dxa', 'bone density', 'osteoporosis'],
            'dermatology': ['skin', 'dermatology', 'mole', 'lesion', 'biopsy'],
            'prostate': ['prostate', 'psa', 'digital rectal'],
            'lung': ['chest ct', 'ldct', 'lung', 'thoracic'],
            'eye': ['eye exam', 'ophthalmology', 'vision', 'retinal'],
            'hearing': ['hearing', 'audiometry', 'audiology']
        }
        
        # Keywords that indicate screening completion
        self.completion_indicators = [
            'normal', 'negative', 'no evidence', 'unremarkable', 'within normal limits',
            'satisfactory', 'adequate', 'complete', 'findings', 'impression',
            'recommendation', 'follow-up', 'next screening'
        ]
    
    def _preprocess_content(self, content: bytes, content_type: Optional[str]) -> Tuple[bytes, Optional[str], Optional[str]]:
        """
        Preprocess document content to detect and handle Epic sandbox artifacts.
        
        Detects and handles:
        1. Base64 encoded content (starts with 'base64;' or data URI patterns)
        2. PDF magic bytes (%PDF) - override file extension regardless of content_type
        3. URL reference pages - garbage content with URLs to actual documents
        4. Image magic bytes (JPEG, PNG) - override file extension
        
        Args:
            content: Raw binary document content
            content_type: MIME type from Epic
            
        Returns:
            Tuple of (processed_content, detected_extension, rejection_reason)
            - processed_content: Decoded/processed bytes
            - detected_extension: Override file extension based on magic bytes (e.g., '.pdf')
            - rejection_reason: If not None, content should be rejected with this reason
        """
        if not content:
            return content, None, "Empty content"
        
        # Try to decode as text to check for text-based artifacts
        try:
            text_content = content.decode('utf-8', errors='ignore')
        except:
            text_content = ""
        
        # Check for base64 encoded content (Epic sometimes returns raw base64)
        if text_content.startswith('base64;') or text_content.startswith('data:'):
            self.logger.info("Detected base64 encoded content, attempting to decode")
            try:
                # Handle 'base64;...' format
                if text_content.startswith('base64;'):
                    b64_data = text_content[7:]  # Remove 'base64;' prefix
                # Handle 'data:mimetype;base64,...' format
                elif ';base64,' in text_content:
                    b64_data = text_content.split(';base64,', 1)[1]
                else:
                    b64_data = text_content
                
                # Decode base64
                decoded_content = base64.b64decode(b64_data)
                self.logger.info(f"Successfully decoded base64 content: {len(decoded_content)} bytes")
                content = decoded_content
                # Re-check the decoded content for magic bytes
                try:
                    text_content = content.decode('utf-8', errors='ignore')
                except:
                    text_content = ""
            except Exception as e:
                self.logger.warning(f"Failed to decode base64 content: {e}")
                return content, None, f"Invalid base64 encoding: {e}"
        
        # Check for URL reference pages (garbage + external URLs)
        # Epic sandbox sometimes returns HTML pages with links to actual documents
        url_patterns = [
            r'https?://[^\s]+ifaxapp\.com[^\s]*',  # iFax URLs
            r'https?://[^\s]+epic[^\s]*/attachment[^\s]*',  # Epic attachment URLs
            r'file:////[0-9]+\.[0-9]+\.[0-9]+\.[0-9]+',  # Internal file:// paths
        ]
        for pattern in url_patterns:
            if re.search(pattern, text_content, re.IGNORECASE):
                # Check if content is mostly URL/garbage (very short with URL)
                clean_text = re.sub(r'https?://[^\s]+', '', text_content).strip()
                clean_text = re.sub(r'file://[^\s]+', '', clean_text).strip()
                # If remaining text is very short or garbage, reject
                if len(clean_text) < 50 or not any(c.isalpha() for c in clean_text[:20]):
                    self.logger.warning(f"Rejecting URL reference page: contains external URL with minimal content")
                    return content, None, "URL reference page - no actual document content"
        
        # Check for PDF magic bytes (%PDF)
        if content[:4] == b'%PDF':
            self.logger.info("Detected PDF magic bytes, overriding to .pdf extension")
            return content, '.pdf', None
        
        # Check for JPEG magic bytes (FFD8FF)
        if content[:3] == b'\xff\xd8\xff':
            self.logger.info("Detected JPEG magic bytes, overriding to .jpg extension")
            return content, '.jpg', None
        
        # Check for PNG magic bytes (89 50 4E 47)
        if content[:4] == b'\x89PNG':
            self.logger.info("Detected PNG magic bytes, overriding to .png extension")
            return content, '.png', None
        
        # Check for TIFF magic bytes (49 49 2A 00 or 4D 4D 00 2A)
        if content[:4] in (b'II*\x00', b'MM\x00*'):
            self.logger.info("Detected TIFF magic bytes, overriding to .tiff extension")
            return content, '.tiff', None
        
        # Check if text content contains raw PDF structure (misidentified as text)
        if '%PDF-' in text_content[:100]:
            self.logger.warning("Detected raw PDF structure in text content - content type mismatch")
            # This is PDF binary being treated as text - convert back to bytes and process as PDF
            return content, '.pdf', None
        
        return content, None, None
    
    def process_document(self, document_content: bytes, document_title: str, content_type: Optional[str] = None, 
                         phi_settings_snapshot=None, fhir_doc_id: Optional[int] = None, 
                         org_id: Optional[int] = None, patient_id: Optional[int] = None) -> Optional[str]:
        """
        Process document content and extract text with screening analysis
        
        Args:
            document_content: Binary document content
            document_title: Document title or filename
            content_type: MIME type (e.g., 'application/pdf', 'image/jpeg')
            phi_settings_snapshot: Optional pre-loaded PHI filter settings for thread-safe
                                  batch processing. If None, settings are queried from DB.
            fhir_doc_id: Optional FHIR document ID for audit logging
            org_id: Optional organization ID for audit logging
            patient_id: Optional patient ID for audit logging
            
        Returns:
            Extracted and processed text or None if processing failed
        """
        try:
            self.logger.info(f"Processing document: {document_title}")
            
            if fhir_doc_id is not None and org_id is not None:
                DocumentAuditLogger.log_processing_started(
                    document_id=fhir_doc_id,
                    document_type='fhir_document',
                    org_id=org_id,
                    patient_id=patient_id
                )
            
            # Preprocess content to handle Epic sandbox artifacts
            # Decodes base64, detects magic bytes, rejects URL reference pages
            processed_content, detected_extension, rejection_reason = self._preprocess_content(document_content, content_type)
            
            if rejection_reason:
                self.logger.warning(f"Document rejected during preprocessing: {rejection_reason}")
                if fhir_doc_id is not None and org_id is not None:
                    DocumentAuditLogger.log_processing_failed(
                        document_id=fhir_doc_id,
                        document_type='fhir_document',
                        org_id=org_id,
                        error_message=f"Content preprocessing failed: {rejection_reason}",
                        patient_id=patient_id
                    )
                return None
            
            # Determine file extension: magic bytes detection > content_type > filename
            if detected_extension:
                file_extension = detected_extension
                self.logger.info(f"Using detected extension from magic bytes: {file_extension}")
            else:
                file_extension = self._get_file_extension_from_content_type(content_type) or self._get_file_extension(document_title)
            
            # Use verified secure deletion for PHI temp files (HIPAA compliance)
            from utils.secure_delete import secure_temp_file, secure_delete_file
            
            with secure_temp_file(suffix=file_extension) as temp_file_path:
                # Write content to secure temp file
                with open(temp_file_path, 'wb') as f:
                    f.write(processed_content)
                
                # Extract text using OCR
                extracted_text, confidence = self._extract_text_from_file(temp_file_path)
                
                # COST CONTROL: Check for oversized document sentinel
                if confidence == -2.0:
                    self.logger.warning(f"COST CONTROL: Document skipped due to page limit: {document_title}")
                    if fhir_doc_id is not None and org_id is not None:
                        DocumentAuditLogger.log_processing_failed(
                            document_id=fhir_doc_id,
                            document_type='fhir_document',
                            org_id=org_id,
                            error_message=f'Document exceeds MAX_DOCUMENT_PAGES limit (cost control)',
                            patient_id=patient_id
                        )
                    return None
                
                if extracted_text:
                    original_length = len(extracted_text)
                    
                    # Apply PHI filtering with counts for audit trail
                    filtered_text, phi_counts = self.phi_filter.filter_phi_with_counts(extracted_text, preloaded_settings=phi_settings_snapshot)
                    
                    if phi_counts and fhir_doc_id is not None and org_id is not None:
                        DocumentAuditLogger.log_phi_redacted(
                            document_id=fhir_doc_id,
                            document_type='fhir_document',
                            org_id=org_id,
                            phi_types_found=phi_counts,
                            original_length=original_length,
                            filtered_length=len(filtered_text),
                            patient_id=patient_id
                        )
                    
                    # Enhance text for screening detection
                    enhanced_text = self._enhance_text_for_screening(filtered_text, document_title)
                    
                    # Binary content guard may reject content
                    if enhanced_text is None:
                        self.logger.warning(f"Document rejected by binary content guard: {document_title}")
                        if fhir_doc_id is not None and org_id is not None:
                            DocumentAuditLogger.log_processing_failed(
                                document_id=fhir_doc_id,
                                document_type='fhir_document',
                                org_id=org_id,
                                error_message='Binary content detected in extracted text (PDF/binary not properly processed)',
                                patient_id=patient_id
                            )
                        return None
                    
                    if fhir_doc_id is not None and org_id is not None:
                        DocumentAuditLogger.log_processing_completed(
                            document_id=fhir_doc_id,
                            document_type='fhir_document',
                            org_id=org_id,
                            confidence=confidence,
                            text_length=len(enhanced_text),
                            processing_method='document_processor',
                            patient_id=patient_id
                        )
                    
                    self.logger.info(f"Successfully processed document with {confidence:.2f} confidence")
                    return enhanced_text
                else:
                    self.logger.warning(f"No text extracted from document: {document_title}")
                    if fhir_doc_id is not None and org_id is not None:
                        DocumentAuditLogger.log_processing_failed(
                            document_id=fhir_doc_id,
                            document_type='fhir_document',
                            org_id=org_id,
                            error_message='No text extracted from document',
                            patient_id=patient_id
                        )
                    return None
            # secure_temp_file context manager handles verified secure deletion on exit
                    
        except Exception as e:
            self.logger.error(f"Error processing document {document_title}: {str(e)}")
            if fhir_doc_id is not None and org_id is not None:
                DocumentAuditLogger.log_processing_failed(
                    document_id=fhir_doc_id,
                    document_type='fhir_document',
                    org_id=org_id,
                    error_message=str(e),
                    patient_id=patient_id
                )
            return None
    
    def analyze_document_for_screenings(self, document_text: str, document_title: str) -> Dict[str, Any]:
        """
        Analyze document text for screening-related content
        
        Args:
            document_text: Extracted document text
            document_title: Document title
            
        Returns:
            Dict with screening analysis results
        """
        try:
            analysis = {
                'screening_types': [],
                'completion_evidence': False,
                'confidence_scores': {},
                'keywords_found': [],
                'document_classification': 'unknown'
            }
            
            text_lower = f"{document_title} {document_text}".lower()
            
            # Check for screening types
            for screening_type, keywords in self.screening_document_types.items():
                max_confidence = 0.0
                found_keywords = []
                
                for keyword in keywords:
                    # Use fuzzy matching for keyword detection
                    confidence = self._calculate_keyword_confidence(keyword, text_lower)
                    
                    if confidence > 0.7:  # High confidence threshold
                        found_keywords.append(keyword)
                        max_confidence = max(max_confidence, confidence)
                
                if max_confidence > 0.7:
                    analysis['screening_types'].append(screening_type)
                    analysis['confidence_scores'][screening_type] = max_confidence
                    analysis['keywords_found'].extend(found_keywords)
            
            # Check for completion indicators
            analysis['completion_evidence'] = self._has_completion_evidence(text_lower)
            
            # Classify document type
            if analysis['screening_types']:
                analysis['document_classification'] = 'screening_report'
            elif any(term in text_lower for term in ['report', 'result', 'findings']):
                analysis['document_classification'] = 'medical_report'
            elif any(term in text_lower for term in ['note', 'visit', 'consultation']):
                analysis['document_classification'] = 'clinical_note'
            
            return analysis
            
        except Exception as e:
            self.logger.error(f"Error analyzing document for screenings: {str(e)}")
            return {'screening_types': [], 'completion_evidence': False, 'confidence_scores': {}, 'keywords_found': [], 'document_classification': 'unknown'}
    
    def extract_screening_dates(self, document_text: str) -> List[Dict[str, Any]]:
        """
        Extract screening dates from document text
        
        Args:
            document_text: Document text to analyze
            
        Returns:
            List of extracted dates with context
        """
        import re
        from datetime import datetime
        
        dates_found = []
        
        # Common date patterns in medical documents
        date_patterns = [
            r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b',  # MM/DD/YYYY or MM-DD-YYYY
            r'\b(\d{1,2}\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{2,4})\b',  # DD Mon YYYY
            r'\b((Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{2,4})\b',  # Mon DD, YYYY
        ]
        
        try:
            for pattern in date_patterns:
                matches = re.finditer(pattern, document_text, re.IGNORECASE)
                
                for match in matches:
                    date_str = match.group(1) if len(match.groups()) > 0 else match.group(0)
                    start_pos = match.start()
                    end_pos = match.end()
                    
                    # Get context around the date
                    context_start = max(0, start_pos - 50)
                    context_end = min(len(document_text), end_pos + 50)
                    context = document_text[context_start:context_end]
                    
                    dates_found.append({
                        'date_string': date_str,
                        'context': context.strip(),
                        'position': start_pos
                    })
            
            return dates_found
            
        except Exception as e:
            self.logger.error(f"Error extracting screening dates: {str(e)}")
            return []
    
    def _detect_content_type_from_magic_bytes(self, file_path: str) -> Optional[str]:
        """
        Detect actual file type from magic bytes (first bytes of file content).
        
        This is crucial for Epic documents that may arrive without proper content_type,
        causing them to be saved with .tmp extension but containing valid PDF/image data.
        
        Returns:
            Detected extension (e.g., '.pdf', '.jpg') or None if unknown
        """
        try:
            with open(file_path, 'rb') as f:
                header = f.read(100)  # Read first 100 bytes for detection
            
            if not header:
                return None
            
            # PDF: %PDF anywhere in first 100 bytes (some PDFs have BOM or whitespace prefix)
            if b'%PDF' in header:
                return '.pdf'
            
            # JPEG: FFD8FF
            if header[:3] == b'\xff\xd8\xff':
                return '.jpg'
            
            # PNG: 89 50 4E 47 0D 0A 1A 0A
            if header[:8] == b'\x89PNG\r\n\x1a\n':
                return '.png'
            
            # TIFF: 49 49 2A 00 (little-endian) or 4D 4D 00 2A (big-endian)
            if header[:4] in (b'II*\x00', b'MM\x00*'):
                return '.tiff'
            
            # BMP: 42 4D
            if header[:2] == b'BM':
                return '.bmp'
            
            # HTML detection
            header_lower = header.lower()
            if b'<!doctype html' in header_lower or b'<html' in header_lower:
                return '.html'
            
            return None
        except Exception as e:
            self.logger.warning(f"Magic byte detection failed for {file_path}: {e}")
            return None
    
    def _extract_text_from_file(self, file_path: str) -> Tuple[Optional[str], float]:
        """
        Extract text from file using appropriate method
        
        Args:
            file_path: Path to document file
            
        Returns:
            Tuple of (extracted_text, confidence_score)
        """
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # For unknown/temp files, detect actual type from magic bytes
            if file_ext in ['.tmp', '']:
                detected_ext = self._detect_content_type_from_magic_bytes(file_path)
                if detected_ext:
                    self.logger.info(f"Detected {detected_ext} from magic bytes for {file_ext} file")
                    file_ext = detected_ext
                else:
                    self.logger.warning(f"Could not detect content type for {file_ext} file, skipping")
                    return None, 0.0
            
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._extract_from_image(file_path)
            elif file_ext in ['.txt']:
                # Plain text file - but verify it's actually text, not misidentified binary
                with open(file_path, 'rb') as f:
                    raw_content = f.read(100)
                
                # Check for binary magic bytes that shouldn't be in text files
                if b'%PDF' in raw_content:
                    self.logger.info("Text file contains PDF data, processing as PDF")
                    return self._extract_from_pdf(file_path)
                if raw_content[:3] == b'\xff\xd8\xff':
                    self.logger.info("Text file contains JPEG data, processing as image")
                    return self._extract_from_image(file_path)
                
                # Actually read as text
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read(), 1.0
            elif file_ext in ['.html', '.htm']:
                # HTML document (Epic clinical notes)
                return self._extract_from_html(file_path)
            elif file_ext == '.docx':
                # Modern Word document
                return self._extract_from_docx(file_path)
            elif file_ext == '.doc':
                # Legacy Word document
                return self._extract_from_doc(file_path)
            else:
                self.logger.warning(f"Unsupported file type: {file_ext}")
                return None, 0.0
                
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None, 0.0
    
    def _get_pdf_page_count(self, pdf_path: str) -> Optional[int]:
        """
        Get page count from PDF without full processing.
        
        Uses PyMuPDF (fitz) if available for efficiency, otherwise
        falls back to pdf2image which is slower but always available.
        
        Returns:
            Page count, or None if unable to determine
        """
        try:
            # Try PyMuPDF first (fast, doesn't render pages)
            try:
                import fitz
                doc = fitz.open(pdf_path)
                page_count = len(doc)
                doc.close()
                return page_count
            except ImportError:
                pass
            
            # Fallback: use pdf2image with just page count (slower but reliable)
            # This actually loads pages, but we just count them
            images = pdf2image.convert_from_path(pdf_path, first_page=1, last_page=1)
            # Use pdfinfo_from_path for page count if available
            try:
                from pdf2image import pdfinfo_from_path
                info = pdfinfo_from_path(pdf_path)
                return info.get('Pages', None)
            except:
                pass
            
            # Last resort: actually convert and count (expensive)
            images = pdf2image.convert_from_path(pdf_path)
            return len(images)
            
        except Exception as e:
            self.logger.warning(f"Could not determine page count for {pdf_path}: {e}")
            return None
    
    def _get_pdf_page_count_from_bytes(self, content: bytes) -> Optional[int]:
        """
        Get page count from PDF bytes without full OCR processing.
        
        COST CONTROL: This is used to check document size before committing
        to expensive OCR processing. Uses PyMuPDF (fitz) for efficiency.
        
        Args:
            content: Raw PDF bytes
            
        Returns:
            Page count, or None if unable to determine
        """
        try:
            # Try PyMuPDF first (fast, works directly from bytes)
            try:
                import fitz
                doc = fitz.open(stream=content, filetype="pdf")
                page_count = len(doc)
                doc.close()
                return page_count
            except ImportError:
                pass
            except Exception as e:
                self.logger.debug(f"PyMuPDF failed to open PDF bytes: {e}")
            
            # Fallback: write to temp file and use pdfinfo
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tf:
                tf.write(content)
                temp_path = tf.name
            
            try:
                from pdf2image import pdfinfo_from_path
                info = pdfinfo_from_path(temp_path)
                return info.get('Pages', None)
            except Exception as e:
                self.logger.debug(f"pdfinfo_from_path failed: {e}")
                return None
            finally:
                try:
                    os.unlink(temp_path)
                except:
                    pass
                    
        except Exception as e:
            self.logger.warning(f"Could not determine page count from PDF bytes: {e}")
            return None
    
    def _extract_from_pdf(self, pdf_path: str, fhir_doc: Optional[FHIRDocument] = None) -> Tuple[Optional[str], float]:
        """Extract text from PDF using hybrid approach with per-page timeout handling
        
        COST OPTIMIZATION (v2): Uses PyMuPDF hybrid extraction:
        1. Try embedded text first (fast, no rendering)
        2. Only render pages that need OCR using get_pixmap() (lazy rendering)
        3. Eliminates pdf2image dependency for most documents
        
        This reduces compute by 50-80% for documents with embedded text.
        
        If a page times out, continues processing remaining pages instead of
        failing the entire document. This matches the graceful degradation
        behavior expected for healthcare document processing.
        
        COST CONTROL: Checks MAX_DOCUMENT_PAGES and skips oversized documents
        to prevent runaway compute costs on large medical records.
        """
        try:
            import fitz
            PYMUPDF_AVAILABLE = True
        except ImportError:
            PYMUPDF_AVAILABLE = False
        
        try:
            max_pages = get_max_document_pages()
            
            # Get page count before processing
            page_count = self._get_pdf_page_count(pdf_path)
            
            # Update FHIRDocument with page count if provided
            if fhir_doc is not None and page_count is not None:
                fhir_doc.page_count = page_count
            
            # COST CONTROL: Skip oversized documents
            if page_count is not None and page_count > max_pages:
                self.logger.warning(
                    f"COST CONTROL: Skipping oversized PDF ({page_count} pages > {max_pages} limit): {pdf_path}"
                )
                if fhir_doc is not None:
                    fhir_doc.skipped_oversized = True
                    fhir_doc.processing_status = 'skipped_oversized'
                    fhir_doc.processing_error = f"Document has {page_count} pages, exceeds limit of {max_pages}"
                # Return special sentinel: -2.0 indicates skipped due to size
                return None, -2.0
            
            with tempfile.TemporaryDirectory() as temp_dir:
                all_text = []
                confidences = []
                timed_out_pages = []
                pages_extracted = 0
                pages_ocred = 0
                
                if PYMUPDF_AVAILABLE:
                    # OPTIMIZED PATH: Use PyMuPDF for hybrid extraction
                    doc = fitz.open(pdf_path)
                    try:
                        num_pages = len(doc)
                        
                        # Update page count if we didn't get it earlier
                        if fhir_doc is not None and fhir_doc.page_count is None:
                            fhir_doc.page_count = num_pages
                        
                        for i in range(num_pages):
                            page = doc[i]
                            page_text = page.get_text("text").strip()
                            
                            # If page has embedded text (>50 chars), use directly (no OCR needed)
                            if len(page_text) >= 50:
                                all_text.append(page_text)
                                confidences.append(1.0)
                                pages_extracted += 1
                            else:
                                # LAZY RENDERING: Only render this page for OCR
                                # Use get_pixmap() instead of pdf2image (faster, no poppler)
                                zoom = 150 / 72  # 150 DPI for good OCR quality
                                matrix = fitz.Matrix(zoom, zoom)
                                pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                                
                                image_path = os.path.join(temp_dir, f"page_{i}.png")
                                pixmap.save(image_path)
                                
                                # Extract text from image (has its own timeout)
                                text, confidence = self._extract_from_image(image_path)
                                
                                if confidence == -1.0:
                                    # Timeout sentinel - page OCR timed out
                                    timed_out_pages.append(i + 1)
                                    self.logger.warning(f"Page {i+1} of {pdf_path} timed out during OCR")
                                    # FALLBACK: Use short embedded text if OCR timed out
                                    if page_text:
                                        all_text.append(page_text)
                                        confidences.append(0.5)  # Lower confidence for short text
                                        pages_extracted += 1
                                elif text:
                                    all_text.append(text)
                                    confidences.append(confidence)
                                    pages_ocred += 1
                                elif page_text:
                                    # OCR returned nothing, but we have some embedded text - use it
                                    all_text.append(page_text)
                                    confidences.append(0.5)  # Lower confidence for short text
                                    pages_extracted += 1
                    finally:
                        doc.close()
                    
                    self.logger.info(
                        f"PDF hybrid extraction: {pages_extracted} pages embedded text, "
                        f"{pages_ocred} pages OCR'd ({pdf_path})"
                    )
                else:
                    # FALLBACK: Use pdf2image if PyMuPDF not available
                    images = pdf2image.convert_from_path(pdf_path)
                    
                    # Update page count if we didn't get it earlier
                    if fhir_doc is not None and fhir_doc.page_count is None:
                        fhir_doc.page_count = len(images)
                    
                    for i, image in enumerate(images):
                        image_path = os.path.join(temp_dir, f"page_{i}.png")
                        image.save(image_path, 'PNG')
                        
                        text, confidence = self._extract_from_image(image_path)
                        
                        if confidence == -1.0:
                            timed_out_pages.append(i + 1)
                            self.logger.warning(f"Page {i+1} of {pdf_path} timed out during OCR")
                        elif text:
                            all_text.append(text)
                            confidences.append(confidence)
                
                if timed_out_pages:
                    self.logger.warning(f"PDF {pdf_path}: {len(timed_out_pages)} pages had no text: {timed_out_pages}")
                
                # Combine all text and calculate average confidence
                combined_text = '\n'.join(all_text) if all_text else None
                avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
                
                return combined_text, avg_confidence
                
        except Exception as e:
            self.logger.error(f"Error processing PDF {pdf_path}: {str(e)}")
            return None, 0.0
    
    def _extract_from_image(self, image_path: str) -> Tuple[Optional[str], float]:
        """Extract text from image using Tesseract OCR with timeout protection
        
        Uses the same OCR_TIMEOUT_SECONDS circuit breaker as manual document processing
        to ensure consistent SLA behavior for both Document and FHIRDocument types.
        """
        timeout_seconds = get_ocr_timeout_seconds()
        
        def _run_tesseract(img_path: str) -> Tuple[Optional[str], float]:
            """Inner function for threaded execution with timeout"""
            try:
                image = Image.open(img_path)
                
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                config = '--oem 3 --psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,()[]{}:;/\\-+=%$@#!?"\' \n\t'
                
                ocr_data = pytesseract.image_to_data(image, config=config, output_type=pytesseract.Output.DICT)
                
                extracted_text = []
                confidences = []
                
                for i in range(len(ocr_data['text'])):
                    confidence = int(ocr_data['conf'][i])
                    text = ocr_data['text'][i].strip()
                    
                    if confidence > 30 and text:
                        extracted_text.append(text)
                        confidences.append(confidence)
                
                if extracted_text:
                    combined_text = ' '.join(extracted_text)
                    avg_confidence = sum(confidences) / len(confidences) / 100.0
                    return combined_text, avg_confidence
                else:
                    return None, 0.0
                    
            except Exception as e:
                self.logger.error(f"Tesseract error for {img_path}: {str(e)}")
                return None, 0.0
        
        try:
            # Use ThreadPoolExecutor with timeout for circuit breaker
            executor = ThreadPoolExecutor(max_workers=1)
            try:
                future = executor.submit(_run_tesseract, image_path)
                done, pending = wait([future], timeout=timeout_seconds)
                
                if done:
                    return future.result()
                else:
                    # Return sentinel value (-1.0) to distinguish timeout from low-confidence OCR
                    self.logger.warning(f"OCR timeout after {timeout_seconds}s for image: {image_path}")
                    return None, -1.0  # Timeout sentinel
            finally:
                executor.shutdown(wait=False, cancel_futures=True)
                
        except Exception as e:
            self.logger.error(f"Error processing image {image_path}: {str(e)}")
            return None, 0.0
    
    def _extract_from_html(self, html_path: str) -> Tuple[Optional[str], float]:
        """Extract text from HTML document (Epic clinical notes)"""
        try:
            import trafilatura
            
            # Read HTML content
            with open(html_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            # Extract clean text using trafilatura
            extracted_text = trafilatura.extract(html_content, include_tables=True, include_comments=False)
            
            if extracted_text:
                self.logger.info(f"Successfully extracted {len(extracted_text)} characters from HTML")
                return extracted_text, 1.0  # HTML extraction is deterministic, high confidence
            else:
                # Fallback: strip HTML tags manually
                from html.parser import HTMLParser
                
                class HTMLTextExtractor(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.text = []
                    
                    def handle_data(self, data):
                        self.text.append(data.strip())
                
                parser = HTMLTextExtractor()
                parser.feed(html_content)
                fallback_text = ' '.join(parser.text)
                
                if fallback_text:
                    self.logger.info(f"Extracted text using fallback HTML parser")
                    return fallback_text, 0.9
                else:
                    return None, 0.0
                
        except Exception as e:
            self.logger.error(f"Error processing HTML {html_path}: {str(e)}")
            return None, 0.0
    
    def _extract_from_docx(self, docx_path: str) -> Tuple[Optional[str], float]:
        """
        Extract text from modern Word documents (.docx)
        
        Uses python-docx library to extract text from paragraphs and tables.
        Since text is directly extracted (not OCR'd), confidence is 1.0.
        
        Args:
            docx_path: Path to the .docx file
            
        Returns:
            Tuple of (extracted_text, confidence_score)
        """
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx library not available for .docx processing")
            return None, 0.0
        
        try:
            doc = DocxDocument(docx_path)  # type: ignore
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            if text_parts:
                combined_text = '\n'.join(text_parts)
                self.logger.info(f"Successfully extracted {len(combined_text)} characters from DOCX")
                return combined_text, 1.0  # Direct text extraction, high confidence
            else:
                self.logger.warning("No text found in DOCX document")
                return None, 0.0
                
        except Exception as e:
            self.logger.error(f"Error processing DOCX {docx_path}: {str(e)}")
            return None, 0.0
    
    def _extract_from_doc(self, doc_path: str) -> Tuple[Optional[str], float]:
        """
        Extract text from legacy Word documents (.doc)
        
        Attempts multiple methods:
        1. antiword utility if available
        2. catdoc utility if available  
        3. Convert to PDF then OCR as fallback
        
        Args:
            doc_path: Path to the .doc file
            
        Returns:
            Tuple of (extracted_text, confidence_score)
        """
        import subprocess
        
        # Method 1: Try antiword
        try:
            result = subprocess.run(
                ['antiword', doc_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
                self.logger.info(f"Successfully extracted {len(text)} characters from DOC using antiword")
                return text, 0.95  # antiword is reliable but may miss some formatting
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        except Exception as e:
            self.logger.warning(f"antiword extraction failed: {str(e)}")
        
        # Method 2: Try catdoc
        try:
            result = subprocess.run(
                ['catdoc', '-w', doc_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
                self.logger.info(f"Successfully extracted {len(text)} characters from DOC using catdoc")
                return text, 0.9  # catdoc may have some formatting issues
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        except Exception as e:
            self.logger.warning(f"catdoc extraction failed: {str(e)}")
        
        # Method 3: Try LibreOffice conversion to PDF then OCR
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', temp_dir, doc_path],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                if result.returncode == 0:
                    # Find the converted PDF
                    pdf_name = os.path.splitext(os.path.basename(doc_path))[0] + '.pdf'
                    pdf_path = os.path.join(temp_dir, pdf_name)
                    
                    if os.path.exists(pdf_path):
                        text, confidence = self._extract_from_pdf(pdf_path)
                        if text:
                            self.logger.info(f"Successfully extracted text from DOC via PDF conversion")
                            return text, confidence * 0.9  # Slightly lower confidence due to conversion
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        except Exception as e:
            self.logger.warning(f"LibreOffice conversion failed: {str(e)}")
        
        self.logger.error(f"Unable to extract text from DOC file: {doc_path}. Install antiword, catdoc, or LibreOffice.")
        return None, 0.0
    
    def _enhance_text_for_screening(self, text: str, document_title: str) -> Optional[str]:
        """
        Enhance extracted text for better screening detection
        
        Args:
            text: Original extracted text
            document_title: Document title for context
            
        Returns:
            Enhanced text with normalized formatting, or None if content is binary/corrupted
        """
        try:
            # BINARY CONTENT GUARD: Reject raw PDF/binary content that wasn't properly extracted
            # This catches cases where PDF content bypassed OCR and was stored as raw text
            binary_signatures = [
                '%PDF-',           # PDF header
                'endstream',       # PDF internal structure
                '/FlateDecode',    # PDF compression marker
                '\x00',            # Null bytes indicate binary
            ]
            for sig in binary_signatures:
                if sig in text[:500]:  # Check first 500 chars
                    self.logger.warning(f"Rejecting binary content in text (found '{sig[:10]}...')")
                    return None
            
            # Start with title for additional context
            enhanced = f"DOCUMENT: {document_title}\n\n{text}"
            
            # Normalize spacing and line breaks
            enhanced = ' '.join(enhanced.split())
            
            # Add section markers for common medical document sections
            section_markers = {
                'FINDINGS:': '\n\nFINDINGS:\n',
                'IMPRESSION:': '\n\nIMPRESSION:\n',
                'RECOMMENDATION:': '\n\nRECOMMENDATION:\n',
                'DIAGNOSIS:': '\n\nDIAGNOSIS:\n',
                'PROCEDURE:': '\n\nPROCEDURE:\n'
            }
            
            for marker, replacement in section_markers.items():
                enhanced = enhanced.replace(marker, replacement)
            
            return enhanced
            
        except Exception as e:
            self.logger.error(f"Error enhancing text: {str(e)}")
            return text
    
    def _calculate_keyword_confidence(self, keyword: str, text: str) -> float:
        """
        Calculate confidence score for keyword match in text
        
        Args:
            keyword: Keyword to search for
            text: Text to search in
            
        Returns:
            Confidence score between 0.0 and 1.0
        """
        try:
            # Direct exact match
            if keyword.lower() in text.lower():
                return 1.0
            
            # Use fuzzy matching for partial matches
            matches = self.fuzzy_engine.fuzzy_match_keywords(text, [keyword], threshold=0.5)
            if matches:
                return matches[0][1]  # Return the confidence score
            
            return 0.0
            
        except Exception as e:
            self.logger.error(f"Error calculating keyword confidence: {str(e)}")
            return 0.0
    
    def _has_completion_evidence(self, text: str) -> bool:
        """
        Check if text contains evidence of screening completion
        
        Args:
            text: Text to analyze
            
        Returns:
            True if completion evidence found
        """
        try:
            for indicator in self.completion_indicators:
                if indicator.lower() in text.lower():
                    return True
            
            # Check for specific completion phrases
            completion_phrases = [
                'screening complete',
                'examination complete',
                'study complete',
                'procedure completed',
                'results available',
                'no further action needed'
            ]
            
            for phrase in completion_phrases:
                if phrase.lower() in text.lower():
                    return True
            
            return False
            
        except Exception as e:
            self.logger.error(f"Error checking completion evidence: {str(e)}")
            return False
    
    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename"""
        if not filename:
            return '.tmp'
        
        _, ext = os.path.splitext(filename)
        return ext if ext else '.tmp'
    
    def _get_file_extension_from_content_type(self, content_type: Optional[str]) -> Optional[str]:
        """
        Convert MIME content type to file extension
        
        Args:
            content_type: MIME type (e.g., 'application/pdf', 'image/jpeg')
            
        Returns:
            File extension (e.g., '.pdf', '.jpg') or None if not recognized
        """
        if not content_type:
            return None
        
        # Common MIME type mappings for medical documents
        mime_to_ext = {
            'application/pdf': '.pdf',
            'image/jpeg': '.jpg',
            'image/jpg': '.jpg',
            'image/png': '.png',
            'image/tiff': '.tiff',
            'image/tif': '.tiff',
            'image/bmp': '.bmp',
            'text/plain': '.txt',
            'text/html': '.html',
            'application/xhtml+xml': '.html',
            'application/msword': '.doc',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx'
        }
        
        # Clean content type (remove charset and other parameters)
        clean_type = content_type.split(';')[0].strip().lower()
        
        return mime_to_ext.get(clean_type)
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """Get document processing statistics"""
        try:
            # This would return processing statistics
            # For now, return a placeholder
            return {
                'documents_processed': 0,
                'screening_documents_identified': 0,
                'average_confidence': 0.0,
                'processing_errors': 0
            }
        except Exception as e:
            self.logger.error(f"Error getting processing statistics: {str(e)}")
            return {}
    
    def process_fhir_documents_batch(self, fhir_document_ids: List[int], max_workers: int = None, 
                                      progress_callback=None) -> Dict[str, Any]:
        """
        Process multiple FHIR documents in parallel using ThreadPoolExecutor.
        
        MATCHES the batch processing pattern from ocr/processor.py for consistency
        between manual Document and FHIRDocument processing.
        
        Uses thread-local sessions for safe database access across threads, matching
        the pattern used in ocr/processor.py.
        
        TIMEOUT HANDLING: Uses the same OCR_TIMEOUT_SECONDS circuit breaker as manual
        document processing to ensure consistent SLA behavior.
        
        Args:
            fhir_document_ids: List of FHIRDocument IDs to process
            max_workers: Maximum number of parallel workers (None = auto-detect)
            progress_callback: Optional callback function(processed, total, current_doc_id)
        
        Returns:
            Dict with results summary including 'timed_out' list if any documents stalled
        """
        from ocr.processor import get_ocr_max_workers
        from app import app, db as app_db
        from sqlalchemy.orm import scoped_session, sessionmaker
        import time
        
        if max_workers is None:
            max_workers = get_ocr_max_workers()
        
        timeout_seconds = get_ocr_timeout_seconds()
        
        results = {
            'total': len(fhir_document_ids),
            'successful': [],
            'failed': [],
            'timed_out': [],
            'skipped_oversized': [],  # COST CONTROL: documents exceeding MAX_DOCUMENT_PAGES
            'start_time': time.time()
        }
        
        if not fhir_document_ids:
            return results
        
        self.logger.info(f"Starting parallel FHIR document processing of {len(fhir_document_ids)} documents with {max_workers} workers")
        
        # Pre-load PHI filter settings snapshot for thread-safe batch processing
        phi_settings_snapshot = self.phi_filter.get_settings_snapshot()
        
        def process_single_fhir_document_with_session(fhir_doc_id: int) -> Tuple[int, str, Optional[str]]:
            """
            Process a single FHIR document with thread-local session for safety.
            
            Returns:
                Tuple of (doc_id, status, error_message)
                status is one of: 'success', 'failed', 'skipped_oversized', 'skipped_already_processed'
            """
            with app.app_context():
                thread_session = scoped_session(sessionmaker(bind=app_db.engine))
                try:
                    fhir_doc = thread_session.query(FHIRDocument).get(fhir_doc_id)
                    if not fhir_doc:
                        return (fhir_doc_id, 'failed', "Document not found")
                    
                    # Skip already processed documents
                    if fhir_doc.is_processed:
                        return (fhir_doc_id, 'skipped_already_processed', None)
                    
                    # Skip already marked as oversized
                    if fhir_doc.skipped_oversized:
                        return (fhir_doc_id, 'skipped_oversized', f"Previously skipped: {fhir_doc.page_count} pages")
                    
                    # Get document content if available
                    if fhir_doc.content_data:
                        content = fhir_doc.content_data
                        title = fhir_doc.title or f"document_{fhir_doc_id}"
                        content_type = fhir_doc.content_type
                        
                        # COST CONTROL: Check page count for PDFs before OCR processing
                        max_pages = get_max_document_pages()
                        if content_type and 'pdf' in content_type.lower():
                            page_count = self._get_pdf_page_count_from_bytes(content)
                            if page_count is not None:
                                fhir_doc.page_count = page_count
                                if page_count > max_pages:
                                    fhir_doc.skipped_oversized = True
                                    fhir_doc.processing_status = 'skipped_oversized'
                                    fhir_doc.processing_error = f"Document has {page_count} pages, exceeds limit of {max_pages}"
                                    thread_session.commit()
                                    self.logger.warning(f"COST CONTROL: Skipping oversized doc {fhir_doc_id} ({page_count} pages > {max_pages} limit)")
                                    return (fhir_doc_id, 'skipped_oversized', f"{page_count} pages exceeds {max_pages} limit")
                        
                        # Process the document (OCR, PHI filtering with pre-loaded settings and audit logging)
                        extracted_text = self.process_document(
                            content, title, content_type, 
                            phi_settings_snapshot=phi_settings_snapshot,
                            fhir_doc_id=fhir_doc_id,
                            org_id=fhir_doc.org_id,
                            patient_id=fhir_doc.patient_id
                        )
                        
                        if extracted_text:
                            # Use mark_processed() helper to properly set all status/audit fields
                            fhir_doc.mark_processed(status='completed', ocr_text=extracted_text)
                            thread_session.commit()
                            return (fhir_doc_id, 'success', None)
                        else:
                            fhir_doc.mark_processed(status='failed', error="No text extracted")
                            thread_session.commit()
                            return (fhir_doc_id, 'failed', "No text extracted")
                    else:
                        return (fhir_doc_id, 'failed', "No content data available")
                        
                except Exception as e:
                    thread_session.rollback()
                    return (fhir_doc_id, 'failed', str(e))
                finally:
                    thread_session.remove()
        
        processed_count = 0
        timed_out_docs = []
        
        executor = ThreadPoolExecutor(max_workers=max_workers)
        try:
            future_to_doc = {executor.submit(process_single_fhir_document_with_session, doc_id): doc_id 
                           for doc_id in fhir_document_ids}
            pending = set(future_to_doc.keys())
            
            while pending:
                done, pending = wait(pending, timeout=timeout_seconds, return_when=FIRST_COMPLETED)
                
                for future in done:
                    doc_id = future_to_doc[future]
                    processed_count += 1
                    
                    try:
                        result_doc_id, status, error = future.result()
                        
                        if status == 'success':
                            results['successful'].append(result_doc_id)
                        elif status == 'skipped_already_processed':
                            results['successful'].append(result_doc_id)  # Count as successful
                        elif status == 'skipped_oversized':
                            results['skipped_oversized'].append({
                                'document_id': result_doc_id,
                                'reason': error or 'Exceeded page limit'
                            })
                        else:
                            results['failed'].append({
                                'document_id': result_doc_id,
                                'error': error or 'Processing failed'
                            })
                            
                    except Exception as e:
                        results['failed'].append({
                            'document_id': doc_id,
                            'error': str(e)
                        })
                    
                    if progress_callback:
                        try:
                            progress_callback(processed_count, len(fhir_document_ids), doc_id)
                        except Exception:
                            pass
                
                # Timeout: mark remaining as stalled
                if not done and pending:
                    for future in pending:
                        doc_id = future_to_doc[future]
                        timed_out_docs.append(doc_id)
                        results['failed'].append({
                            'document_id': doc_id,
                            'error': f'OCR timeout after {timeout_seconds}s (circuit breaker)'
                        })
                        self.logger.warning(f"FHIR Document {doc_id} exceeded timeout of {timeout_seconds}s")
                    break
        finally:
            executor.shutdown(wait=False, cancel_futures=True)
        
        if timed_out_docs:
            results['timed_out'] = timed_out_docs
            self.logger.warning(f"FHIR document processing: {len(timed_out_docs)} documents timed out")
        
        results['end_time'] = time.time()
        results['duration_seconds'] = results['end_time'] - results['start_time']
        results['docs_per_second'] = len(fhir_document_ids) / results['duration_seconds'] if results['duration_seconds'] > 0 else 0
        
        skipped_count = len(results['skipped_oversized'])
        log_msg = f"Parallel FHIR OCR complete: {len(results['successful'])}/{len(fhir_document_ids)} successful"
        if skipped_count > 0:
            log_msg += f", {skipped_count} skipped (COST CONTROL: exceeded page limit)"
        log_msg += f", {results['docs_per_second']:.2f} docs/sec"
        self.logger.info(log_msg)
        
        return results